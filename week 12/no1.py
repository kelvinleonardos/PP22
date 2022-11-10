from docx import Document
from docx.shared import Cm
docx = Document()
docx.add_heading('BIODATA')

table = docx.add_table(rows=7, cols=3)

for cell in table.columns[1].cells:
    cell.width = Cm(0.5)

row = table.rows[0]
row.cells[0].text = 'Nama'
row.cells[1].text = ':'
row.cells[2].text = 'Kelvin'

row = table.rows[1]
row.cells[0].text = 'Tempat Lahir'
row.cells[1].text = ':'
row.cells[2].text = 'Jakarta'

row = table.rows[2]
row.cells[0].text = 'Tanggal Lahir'
row.cells[1].text = ':'
row.cells[2].text = '17 Mei 2004'

row = table.rows[3]
row.cells[0].text = 'Alamat'
row.cells[1].text = ':'
row.cells[2].text = 'Makassar'

row = table.rows[4]
row.cells[0].text = 'Fakultas'
row.cells[1].text = ':'
row.cells[2].text = 'FMIPA'

row = table.rows[5]
row.cells[0].text = 'Departemen'
row.cells[1].text = ':'
row.cells[2].text = 'Matematika'

row = table.rows[6]
row.cells[0].text = 'Prodi'
row.cells[1].text = ':'
row.cells[2].text = 'Sistem Informasi'

docx.save('BIODATA.docx')